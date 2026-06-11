from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import os
import re
import subprocess
import tempfile
import uuid

import numpy as np
import requests
from docx import Document
from pypdf import PdfReader


ALLOWED_EXTENSIONS = {".pdf", ".docx"}
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 160
MAX_LOGICAL_CHUNK_CHARS = 2200
DEFAULT_TOP_K = 2
MAX_MEMORY_MESSAGES = 3
MIN_GOOD_EXTRACTION_CHARS = 350
MIN_READABLE_WORD_RATIO = 0.55
OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "llama3.2:3b"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
PDF_EXTRACTION_DEBUG = os.getenv("PDF_EXTRACTION_DEBUG", "0").strip().lower() in {"1", "true", "yes", "on"}

RESUME_SECTIONS = {
    "Contact Information": {
        "contact",
        "contact information",
        "personal information",
        "personal details",
        "contact details",
    },

    "Summary": {
        "summary",
        "professional summary",
        "career summary",
        "profile",
        "professional profile",
        "objective",
        "career objective",
        "about me",
        "about",
    },

    "Education": {
        "education",
        "academic background",
        "academics",
        "academic qualifications",
        "academic qualification",
        "qualification",
        "qualifications",
        "educational qualification",
        "educational qualifications",
    },

    "Skills": {
        "skills",
        "technical skills",
        "technical skill",
        "core skills",
        "key skills",
        "competencies",
        "technical competencies",
        "technical expertise",
        "expertise",
        "technologies",
        "tech stack",
        "programming languages",
        "tools",
        "frameworks",
    },

    "Experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "employment history",
        "work history",
        "internship",
        "internships",
        "internship experience",
        "professional experience",
    },

    "Projects": {
        "projects",
        "project",
        "academic projects",
        "personal projects",
        "professional projects",
        "selected projects",
        "key projects",
        "major projects",
        "featured projects",
        "projects undertaken",
        "project experience",
        "projects and internships",
        "projects internships",
    },

    "Certifications": {
        "certifications",
        "certification",
        "certificates",
        "courses",
        "training",
        "licenses",
        "licenses and certifications",
    },

    "Achievements": {
        "achievements",
        "achievement",
        "awards",
        "honors",
        "honours",
        "accomplishments",
    },

    "Languages": {
        "languages",
        "language",
        "known languages",
        "language proficiency",
    },
}

SUMMARY_QUERY_TERMS = {
    "summarize",
    "summary",
    "overview",
    "profile",
    "recruiter summary",
    "career highlights",
    "highlight",
    "highlights",
}

SUMMARY_SECTION_PRIORITY = [
    "Summary",
    "Experience",
    "Projects",
    "Skills",
    "Education",
    "Certifications",
    "Achievements",
]


@dataclass
class RagChunk:
    chunk_number: int
    text: str
    section: str
    document_id: str
    document_name: str

    @property
    def size(self):
        return len(self.text)


_embedding_model = None
_faiss = None
_faiss_index = None
_vector_chunks = []
_vector_ids = set()


def create_vector_store():
    return {
        "index": None,
        "chunks": [],
        "ids": set(),
    }


def reset_vector_store(vector_store=None):
    global _faiss_index, _vector_chunks, _vector_ids
    if vector_store is None:
        _faiss_index = None
        _vector_chunks = []
        _vector_ids = set()
    else:
        vector_store["index"] = None
        vector_store["chunks"] = []
        vector_store["ids"] = set()


def validate_file_type(file_path):
    extension = Path(file_path).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{extension}'. Allowed file types: {allowed}.")
    return extension


def extract_text(file_path):
    extension = validate_file_type(file_path)

    if extension == ".pdf":
        return _extract_pdf_text(file_path)

    if extension == ".docx":
        return _extract_docx_text(file_path)

    return _extract_doc_text(file_path)


def _extract_pdf_text(file_path):
    extraction_attempts = [
        ("pymupdf", _extract_pdf_text_pymupdf),
        ("pdfplumber", _extract_pdf_text_pdfplumber),
        ("pypdf", _extract_pdf_text_pypdf),
    ]
    best_result = None

    for method, extractor in extraction_attempts:
        try:
            result = extractor(file_path)
        except Exception as exc:
            print(f"PDF extraction via {method} failed: {exc}")
            continue

        result["extraction_method"] = method
        result["quality"] = evaluate_extraction_quality(result["text"])
        best_result = choose_better_extraction(best_result, result)

        if is_excellent_extraction(result["quality"]):
            return result

    if best_result and best_result["quality"]["is_poor"]:
        for ocr_method, ocr_extractor in (
            ("easyocr", _extract_pdf_text_easyocr),
            ("paddleocr", _extract_pdf_text_paddleocr),
        ):
            try:
                ocr_result = ocr_extractor(file_path)
                ocr_result["extraction_method"] = ocr_method
                ocr_result["quality"] = evaluate_extraction_quality(ocr_result["text"])
                best_result = choose_better_extraction(best_result, ocr_result)

                if not ocr_result["quality"]["is_poor"]:
                    break
            except Exception as exc:
                print(f"OCR fallback via {ocr_method} skipped or failed: {exc}")

    if best_result:
        return best_result

    return {
        "text": "",
        "page_count": 0,
        "extraction_method": "none",
        "quality": evaluate_extraction_quality(""),
    }


def _extract_pdf_text_pymupdf(file_path):
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is not installed.") from exc

    pages = []

    with fitz.open(file_path) as document:
        for page_num, page in enumerate(document, start=1):
            raw_blocks = page.get_text("blocks")
            text_blocks = normalize_pdf_blocks(raw_blocks)
            layout = detect_pdf_layout(text_blocks, page.rect.width)
            ordered_blocks = order_pdf_blocks_for_reading(text_blocks, layout, page.rect.width)
            debug_pdf_blocks(page_num, text_blocks, layout)
            page_text = "\n".join(block["text"] for block in ordered_blocks if block["text"])

            if page_text:
                pages.append(f"\n[Page {page_num}]\n{page_text}")

        return {
            "text": "\n".join(pages).strip(),
            "page_count": document.page_count,
        }


def normalize_pdf_blocks(raw_blocks):
    blocks = []
    for block in raw_blocks:
        if len(block) < 5:
            continue
        x0, y0, x1, y1, text = block[:5]
        text = re.sub(r"\n{3,}", "\n\n", str(text)).strip()
        if not text:
            continue
        blocks.append(
            {
                "x0": float(x0),
                "y0": float(y0),
                "x1": float(x1),
                "y1": float(y1),
                "text": text,
                "width": float(x1) - float(x0),
                "height": float(y1) - float(y0),
                "center_x": (float(x0) + float(x1)) / 2,
            }
        )
    return blocks


def detect_pdf_layout(blocks, page_width):
    if len(blocks) < 4:
        return {
            "type": "single",
            "split_x": page_width / 2,
            "header_bottom": 0,
        }

    full_width_threshold = page_width * 0.62
    top_limit = min(block["y0"] for block in blocks) + 170
    header_blocks = [
        block
        for block in blocks
        if block["y0"] <= top_limit and block["width"] >= full_width_threshold
    ]
    header_bottom = max((block["y1"] for block in header_blocks), default=0)
    body_blocks = [
        block
        for block in blocks
        if block["y0"] > header_bottom + 4 and block["width"] < page_width * 0.88
    ]

    if len(body_blocks) < 4:
        return {
            "type": "single",
            "split_x": page_width / 2,
            "header_bottom": header_bottom,
        }

    centers = sorted(block["center_x"] for block in body_blocks)
    gaps = [
        (centers[index + 1] - centers[index], centers[index], centers[index + 1])
        for index in range(len(centers) - 1)
    ]
    largest_gap = max(gaps, default=(0, 0, 0), key=lambda item: item[0])
    split_x = (largest_gap[1] + largest_gap[2]) / 2 if largest_gap[0] else page_width / 2
    left_count = sum(1 for block in body_blocks if block["center_x"] < split_x)
    right_count = len(body_blocks) - left_count
    is_multi = (
        largest_gap[0] >= page_width * 0.08
        and left_count >= 2
        and right_count >= 2
    )

    return {
        "type": "multi" if is_multi else "single",
        "split_x": split_x,
        "header_bottom": header_bottom,
        "left_count": left_count,
        "right_count": right_count,
        "largest_gap": largest_gap[0],
    }


def order_pdf_blocks_for_reading(blocks, layout, page_width):
    if not blocks:
        return []

    header_blocks = [
        block
        for block in blocks
        if layout["header_bottom"] and block["y1"] <= layout["header_bottom"] + 2
    ]
    body_blocks = [
        block
        for block in blocks
        if block not in header_blocks
    ]
    ordered_headers = sort_blocks_top_left(header_blocks)

    if layout["type"] != "multi":
        return ordered_headers + sort_blocks_top_left(body_blocks)

    split_x = layout["split_x"]
    full_width_body = [
        block
        for block in body_blocks
        if block["width"] >= page_width * 0.72
    ]
    column_blocks = [
        block
        for block in body_blocks
        if block not in full_width_body
    ]
    left_blocks = [block for block in column_blocks if block["center_x"] < split_x]
    right_blocks = [block for block in column_blocks if block["center_x"] >= split_x]

    ordered_body = []
    ordered_body.extend(sort_blocks_top_left(full_width_body))
    ordered_body.extend(sort_blocks_top_left(left_blocks))
    ordered_body.extend(sort_blocks_top_left(right_blocks))
    return ordered_headers + ordered_body


def sort_blocks_top_left(blocks):
    return sorted(blocks, key=lambda block: (round(block["y0"], 1), round(block["x0"], 1)))


def debug_pdf_blocks(page_num, blocks, layout):
    if not PDF_EXTRACTION_DEBUG:
        return

    print("\n" + "=" * 50)
    print(
        f"PyMuPDF page {page_num} layout={layout.get('type')} "
        f"split_x={layout.get('split_x'):.1f} header_bottom={layout.get('header_bottom'):.1f}"
    )
    for index, block in enumerate(sort_blocks_top_left(blocks), start=1):
        preview = re.sub(r"\s+", " ", block["text"])[:90]
        print(
            f"Block {index}: "
            f"x0={block['x0']:.1f}, y0={block['y0']:.1f}, "
            f"x1={block['x1']:.1f}, y1={block['y1']:.1f}, "
            f"w={block['width']:.1f} | {preview}"
        )
    print("=" * 50)


def _extract_pdf_text_pdfplumber(file_path):
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber is not installed.") from exc

    pages = []
    with pdfplumber.open(file_path) as document:
        for index, page in enumerate(document.pages, start=1):
            page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            if page_text:
                pages.append(f"\n[Page {index}]\n{page_text.strip()}")

        return {
            "text": "\n".join(pages).strip(),
            "page_count": len(document.pages),
        }


def _extract_pdf_text_pypdf(file_path):
    reader = PdfReader(file_path)
    pages = []

    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text()
        if page_text:
            pages.append(f"\n[Page {index}]\n{page_text.strip()}")

    return {
        "text": "\n".join(pages).strip(),
        "page_count": len(reader.pages),
    }


def _extract_pdf_text_easyocr(file_path):
    try:
        import easyocr
        import fitz
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("EasyOCR fallback requires easyocr, PyMuPDF, and Pillow.") from exc

    reader = easyocr.Reader(["en"], gpu=False)
    pages = []

    with fitz.open(file_path) as document:
        for index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            results = reader.readtext(np.array(image), detail=0, paragraph=True)
            page_text = "\n".join(results).strip()
            if page_text:
                pages.append(f"\n[Page {index}]\n{page_text}")

        return {
            "text": "\n".join(pages).strip(),
            "page_count": document.page_count,
        }


def _extract_pdf_text_paddleocr(file_path):
    try:
        import fitz
        from paddleocr import PaddleOCR
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError("PaddleOCR fallback requires paddleocr, PyMuPDF, and Pillow.") from exc

    ocr = PaddleOCR(use_angle_cls=True, lang="en", show_log=False)
    pages = []

    with fitz.open(file_path) as document:
        for index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            results = ocr.ocr(np.array(image), cls=True)
            lines = []

            for page_result in results or []:
                for item in page_result or []:
                    if item and len(item) >= 2 and item[1]:
                        lines.append(str(item[1][0]))

            page_text = "\n".join(lines).strip()
            if page_text:
                pages.append(f"\n[Page {index}]\n{page_text}")

        return {
            "text": "\n".join(pages).strip(),
            "page_count": document.page_count,
        }


def evaluate_extraction_quality(text):
    clean_text = text or ""
    words = re.findall(r"[A-Za-z][A-Za-z.+#'-]{1,}", clean_text)
    tokens = re.findall(r"\S+", clean_text)
    readable_ratio = len(words) / max(1, len(tokens))
    has_email = bool(re.search(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+", clean_text))
    has_phone = bool(re.search(r"\+?\d[\d\s().-]{8,}\d", clean_text))
    has_resume_section = any(
        re.search(rf"\b{re.escape(alias)}\b", clean_text, flags=re.IGNORECASE)
        for aliases in RESUME_SECTIONS.values()
        for alias in aliases
    )
    scrambled_score = calculate_scrambled_score(clean_text)
    is_poor = (
        len(clean_text.strip()) < MIN_GOOD_EXTRACTION_CHARS
        or readable_ratio < MIN_READABLE_WORD_RATIO
        or scrambled_score > 0.18
        or not (has_email or has_phone or has_resume_section)
    )

    return {
        "char_count": len(clean_text.strip()),
        "readable_word_ratio": round(readable_ratio, 3),
        "has_email": has_email,
        "has_phone": has_phone,
        "has_resume_section": has_resume_section,
        "scrambled_score": round(scrambled_score, 3),
        "is_poor": is_poor,
    }


def calculate_scrambled_score(text):
    if not text:
        return 1.0
    odd_chars = re.findall(r"[^\w\s.,;:!?@#%&()\-+/|']", text)
    return len(odd_chars) / max(1, len(text))


def choose_better_extraction(current, candidate):
    if current is None:
        return candidate

    current_quality = current.get("quality") or evaluate_extraction_quality(current.get("text", ""))
    candidate_quality = candidate.get("quality") or evaluate_extraction_quality(candidate.get("text", ""))
    if current_quality["is_poor"] != candidate_quality["is_poor"]:
        return candidate if not candidate_quality["is_poor"] else current

    current_score = extraction_quality_score(current_quality)
    candidate_score = extraction_quality_score(candidate_quality)
    return candidate if candidate_score > current_score else current


def is_excellent_extraction(quality):
    return (
        not quality["is_poor"]
        and quality["char_count"] >= MIN_GOOD_EXTRACTION_CHARS
        and quality["readable_word_ratio"] >= 0.82
        and quality["scrambled_score"] <= 0.04
        and quality["has_resume_section"]
        and (quality["has_email"] or quality["has_phone"])
    )


def extraction_quality_score(quality):
    score = quality["char_count"]
    score += 250 if quality["has_email"] else 0
    score += 200 if quality["has_phone"] else 0
    score += 250 if quality["has_resume_section"] else 0
    score += int(quality["readable_word_ratio"] * 500)
    score -= int(quality["scrambled_score"] * 1000)
    return score


def _extract_docx_text(file_path):
    document = Document(file_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n\n".join(paragraphs)
    return {
        "text": text,
        "page_count": None,
        "extraction_method": "docx",
        "quality": evaluate_extraction_quality(text),
    }


def _extract_doc_text(file_path):
    converted_path = _convert_doc_to_docx(file_path)
    try:
        result = _extract_docx_text(converted_path)
        result["page_count"] = None
        return result
    finally:
        try:
            Path(converted_path).unlink(missing_ok=True)
        except OSError:
            pass


def _convert_doc_to_docx(file_path):
    source_path = Path(file_path)

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir_path = Path(temp_dir)
        command = [
            "soffice",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(temp_dir_path),
            str(source_path),
        ]

        try:
            completed = subprocess.run(command, capture_output=True, text=True, check=False)
        except FileNotFoundError as exc:
            raise RuntimeError(
                "Legacy DOC files require LibreOffice to be installed and available as 'soffice'. "
                "Please upload a PDF or DOCX file, or install LibreOffice to process DOC files."
            ) from exc

        converted_path = temp_dir_path / f"{source_path.stem}.docx"
        if completed.returncode != 0 or not converted_path.exists():
            raise RuntimeError(
                "Could not convert the DOC file to DOCX. "
                f"LibreOffice output: {completed.stderr or completed.stdout}"
            )

        persistent_path = Path(tempfile.gettempdir()) / converted_path.name
        persistent_path.write_bytes(converted_path.read_bytes())
        return persistent_path


def create_chunks(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP, document_id=None, document_name="Document"):
    if not text:
        return []

    document_id = document_id or uuid.uuid4().hex
    chunks = []
    chunk_number = 1
    sections = split_resume_sections(text)

    for section in sections:
        if section["section"] == "Projects":
            logical_parts = split_project_section(section)
        else:
            logical_parts = split_large_section_by_paragraph(section)

        for part in logical_parts:
            content = part["content"].strip()
            if not content:
                continue

            chunk_id = f"{document_id}-{chunk_number}"
            chunks.append(
                build_chunk(
                    chunk_id=chunk_id,
                    chunk_number=chunk_number,
                    document_id=document_id,
                    document_name=document_name,
                    section=part["section"],
                    title=part["title"],
                    page=part["page"],
                    content=content,
                )
            )
            chunk_number += 1

    return chunks


def build_chunk(chunk_id, chunk_number, document_id, document_name, section, title, page, content):
    return {
        "chunk_id": chunk_id,
        "chunk_number": chunk_number,
        "section": section,
        "title": title or section,
        "page": page,
        "content": content,
        "text": content,
        "document_id": document_id,
        "document_name": document_name,
        "size": len(content),
    }


def split_resume_sections(text):
    sections = []
    current_section = "Contact Information"
    current_title = "Contact Information"
    current_page = None
    lines = []
    seen_heading = False

    def flush():
        content = "\n".join(line for line, _ in lines).strip()
        if not content:
            return

        page = first_page(lines) or current_page
        sections.append(
            {
                "section": current_section,
                "title": current_title,
                "page": page,
                "lines": list(lines),
                "content": content,
            }
        )

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        page_match = re.match(r"^\[Page\s+(\d+)\]$", line.strip(), flags=re.IGNORECASE)
        if page_match:
            current_page = int(page_match.group(1))
            continue

        detected_section = detect_resume_section(line)
        if detected_section:
            flush()
            lines = []
            current_section = detected_section
            current_title = detected_section
            seen_heading = True
            remainder = section_heading_remainder(line, detected_section)
            if remainder:
                lines.append((remainder, current_page))
            continue

        if line.strip() or seen_heading or current_section == "Contact Information":
            lines.append((line, current_page))

    flush()

    if sections:
        return sections

    return [
        {
            "section": "Summary",
            "title": "Summary",
            "page": None,
            "lines": [(text, None)],
            "content": text.strip(),
        }
    ]


def detect_resume_section(line):
    clean = normalize_heading(line)
    if not clean or len(clean) > 80:
        return None

    for section, aliases in RESUME_SECTIONS.items():
        section_aliases = aliases | {section.lower()}
        if clean in section_aliases:
            return section
        for alias in section_aliases:
            alias_clean = normalize_heading(alias)
            if clean.startswith(f"{alias_clean} ") and len(clean.split()) <= len(alias_clean.split()) + 3:
                return section
            if re.match(rf"^\s*{re.escape(alias)}\s*[:|-]", line, flags=re.IGNORECASE):
                return section
    return None


def section_heading_remainder(line, section):
    clean_line = line.strip()
    if ":" not in clean_line:
        return ""

    heading, remainder = clean_line.split(":", 1)
    if detect_resume_section(heading) == section:
        return remainder.strip()
    return ""


def normalize_heading(line):
    clean = re.sub(r"\s+", " ", line).strip().lower()
    clean = collapse_spaced_heading(clean)
    clean = clean.strip(":-|")
    clean = re.sub(r"^[#*\-\u2022\s]+", "", clean)
    clean = re.sub(r"[^a-z\s/&]", "", clean)
    clean = clean.replace("&", " and ")
    clean = clean.replace("/", " ")
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean


def collapse_spaced_heading(value):
    tokens = value.split()
    if len(tokens) >= 3 and all(len(token.strip(":-|")) == 1 for token in tokens):
        return "".join(token.strip(":-|") for token in tokens)
    return value


def split_large_section_by_paragraph(section):
    content = section["content"].strip()
    if len(content) <= MAX_LOGICAL_CHUNK_CHARS:
        return [section]

    paragraphs = paragraphs_from_lines(section["lines"])
    if not paragraphs:
        return [section]

    chunks = []
    current_lines = []

    for paragraph in paragraphs:
        paragraph_text = lines_to_text(paragraph)
        current_text = lines_to_text(current_lines)
        would_exceed = len(current_text) + len(paragraph_text) + 2 > MAX_LOGICAL_CHUNK_CHARS

        if current_lines and would_exceed:
            chunks.append(section_part(section, current_lines, section["title"]))
            current_lines = []

        current_lines.extend(paragraph)

    if current_lines:
        chunks.append(section_part(section, current_lines, section["title"]))

    return chunks


def split_project_section(section):
    project_blocks = split_project_blocks(section["lines"])
    if not project_blocks:
        return split_large_section_by_paragraph(section)

    chunks = []
    for block in project_blocks:
        title = infer_project_title(block) or "Project"
        block_text = lines_to_text(block)

        if len(block_text) <= MAX_LOGICAL_CHUNK_CHARS:
            chunks.append(section_part(section, block, title))
            continue

        for paragraph in paragraphs_from_lines(block):
            paragraph_text = lines_to_text(paragraph)
            if paragraph_text:
                chunks.append(section_part(section, paragraph, title))

    return chunks


def split_project_blocks(lines):
    blocks = []
    current = []
    previous_blank = True

    for index, item in enumerate(lines):
        line = item[0].strip()
        if not line:
            if current:
                current.append(item)
            previous_blank = True
            continue

        starts_new_project = current and is_project_title_line(line) and (
            previous_blank or is_distinct_project_name(line)
        )
        if starts_new_project:
            blocks.append(trim_blank_lines(current))
            current = [item]
        else:
            current.append(item)
        previous_blank = False

    if current:
        blocks.append(trim_blank_lines(current))

    if len(blocks) == 1:
        paragraphs = paragraphs_from_lines(lines)
        title_like_count = sum(1 for paragraph in paragraphs if paragraph and is_project_title_line(paragraph[0][0].strip()))
        if title_like_count > 1:
            return paragraphs

    return [block for block in blocks if lines_to_text(block)]


def is_project_title_line(line):
    clean = strip_list_marker(line).strip()
    if not clean:
        return False
    if re.match(
        r"^(tech stack|technologies|tools|role|duration|description|github|link|live|demo|features|responsibilities|impact)\s*:",
        clean,
        flags=re.IGNORECASE,
    ):
        return False
    if detect_resume_section(clean):
        return False
    if len(clean) > 120:
        return False
    if clean.endswith("."):
        return False
    if re.search(r"https?://|@|\b\d{10}\b", clean, flags=re.IGNORECASE):
        return False
    return len(clean.split()) <= 14


def infer_project_title(lines):
    for line, _ in lines:
        clean = strip_list_marker(line).strip(" :-|")
        if clean:
            return clean[:120]
    text = lines_to_text(lines)
    return text[:80] if text else None


def is_distinct_project_name(line):
    clean = strip_list_marker(line).strip(" :-|")
    if not clean or len(clean.split()) > 3:
        return False
    return bool(re.search(r"[a-z][A-Z]", clean)) or bool(re.match(r"^[A-Z][A-Za-z0-9_-]{3,}$", clean))


def paragraphs_from_lines(lines):
    paragraphs = []
    current = []

    for item in lines:
        line = item[0]
        if line.strip():
            current.append(item)
        elif current:
            paragraphs.append(current)
            current = []

    if current:
        paragraphs.append(current)

    return paragraphs


def section_part(section, lines, title):
    return {
        "section": section["section"],
        "title": title,
        "page": first_page(lines) or section.get("page"),
        "content": lines_to_text(lines),
        "lines": lines,
    }


def first_page(lines):
    for _, page in lines:
        if page is not None:
            return page
    return None


def lines_to_text(lines):
    return "\n".join(line for line, _ in lines).strip()


def trim_blank_lines(lines):
    trimmed = list(lines)
    while trimmed and not trimmed[0][0].strip():
        trimmed.pop(0)
    while trimmed and not trimmed[-1][0].strip():
        trimmed.pop()
    return trimmed


def is_bullet_line(line):
    return bool(re.match(r"^(\*|-|\u2022|\d+[\).\s])\s*", line.strip()))


def strip_list_marker(line):
    return re.sub(r"^(\*|-|\u2022|\d+[\).\s])\s*", "", line.strip())


def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        try:
            from sentence_transformers import SentenceTransformer
        except ImportError as exc:
            raise RuntimeError(
                "sentence-transformers is required for semantic retrieval. "
                "Install dependencies with 'pip install -r requirements.txt'."
            ) from exc
        _embedding_model = SentenceTransformer(EMBEDDING_MODEL)
    return _embedding_model


def get_faiss():
    global _faiss
    if _faiss is None:
        try:
            import faiss
        except ImportError as exc:
            raise RuntimeError(
                "faiss-cpu is required for vector search. "
                "Install dependencies with 'pip install -r requirements.txt'."
            ) from exc
        _faiss = faiss
    return _faiss


def get_faiss_index(embedding_dimension, vector_store=None):
    global _faiss_index
    if vector_store is not None:
        if vector_store.get("index") is None:
            faiss = get_faiss()
            vector_store["index"] = faiss.IndexFlatIP(embedding_dimension)
        return vector_store["index"]

    if _faiss_index is None:
        faiss = get_faiss()
        _faiss_index = faiss.IndexFlatIP(embedding_dimension)
    return _faiss_index


def build_vector_index(document, vector_store=None):
    global _vector_chunks

    chunks = document["chunks"]
    if not chunks:
        return

    model = get_embedding_model()
    texts = [chunk["text"] for chunk in chunks]
    embeddings = model.encode(texts, normalize_embeddings=True)
    embeddings = np.asarray(embeddings, dtype="float32")
    index = get_faiss_index(embeddings.shape[1], vector_store=vector_store)

    new_embeddings = []
    new_chunks = []
    vector_ids = vector_store["ids"] if vector_store is not None else _vector_ids

    for chunk, embedding in zip(chunks, embeddings):
        vector_id = f"{document['document_id']}-{chunk['chunk_number']}"
        if vector_id in vector_ids:
            continue

        vector_ids.add(vector_id)
        new_embeddings.append(embedding)
        new_chunks.append(
            {
                "chunk_id": chunk["chunk_id"],
                "text": chunk["text"],
                "content": chunk["content"],
                "document_id": chunk["document_id"],
                "document_name": chunk["document_name"],
                "chunk_number": chunk["chunk_number"],
                "section": chunk["section"],
                "title": chunk["title"],
                "page": chunk["page"],
                "size": chunk["size"],
            }
        )

    if new_embeddings:
        index.add(np.asarray(new_embeddings, dtype="float32"))
        if vector_store is None:
            _vector_chunks.extend(new_chunks)
        else:
            vector_store["chunks"].extend(new_chunks)


def process_document(file_path, vector_store=None, reset_store=True, document_name=None):
    if reset_store:
        reset_vector_store(vector_store)

    path = Path(file_path)
    extraction = extract_text(path)
    text = extraction["text"]
    print("\n" + "=" * 80)
    print("EXTRACTED TEXT")
    print("=" * 80)
    print(text[:5000])   # first 5000 chars
    print("=" * 80)
    document_id = uuid.uuid4().hex
    display_name = document_name or path.name
    chunks = create_chunks(text, document_id=document_id, document_name=display_name)
    log_generated_chunks(display_name, text, chunks)

    document = {
        "document_id": document_id,
        "name": display_name,
        "text": text,
        "preview": text[:2500],
        "chunks": chunks,
        "page_count": extraction["page_count"],
        "character_count": len(text),
        "chunk_count": len(chunks),
        "indexed_at": datetime.now().isoformat(timespec="seconds"),
        "extraction_method": extraction.get("extraction_method", "docx" if path.suffix.lower() == ".docx" else "unknown"),
        "extraction_quality": extraction.get("quality", evaluate_extraction_quality(text)),
    }
    build_vector_index(document, vector_store=vector_store)
    return document


def retrieve_relevant_chunks(question, document_or_documents, top_k=None, vector_store=None):
    documents = _as_document_list(document_or_documents)
    document_ids = [document["document_id"] for document in documents]
    if not document_ids:
        return {"chunks": [], "context": "", "context_size": 0, "chunk_count": 0}

    document_chunks = [
        chunk
        for document in documents
        for chunk in document.get("chunks", [])
        if chunk.get("document_id") in document_ids
    ]
    vector_chunks = vector_store["chunks"] if vector_store is not None else _vector_chunks

    direct_retrieval = retrieve_by_resume_intent(question, document_chunks)
    if direct_retrieval is not None:
        return build_retrieval_payload(
            direct_retrieval["chunks"],
            strategy=direct_retrieval["strategy"],
            query_type=direct_retrieval["query_type"],
        )

    model = get_embedding_model()
    query_embedding = model.encode([question], normalize_embeddings=True)
    query_embedding = np.asarray(query_embedding, dtype="float32")
    index = get_faiss_index(query_embedding.shape[1], vector_store=vector_store)

    if index.ntotal == 0:
        return {"chunks": [], "context": "", "context_size": 0, "chunk_count": 0}

    search_count = index.ntotal
    scores, indexes = index.search(query_embedding, search_count)

    retrieved = []
    semantic_top_k = top_k or DEFAULT_TOP_K
    for score, chunk_index in zip(scores[0], indexes[0]):
        if chunk_index < 0 or chunk_index >= len(vector_chunks):
            continue

        metadata = vector_chunks[chunk_index]
        if metadata["document_id"] not in document_ids:
            continue

        retrieved.append(
            {
                "text": metadata["text"],
                "content": metadata.get("content", metadata["text"]),
                "chunk_id": metadata.get("chunk_id"),
                "document_id": metadata.get("document_id"),
                "document_name": metadata.get("document_name"),
                "chunk_number": metadata.get("chunk_number"),
                "section": metadata.get("section"),
                "title": metadata.get("title"),
                "page": metadata.get("page"),
                "size": metadata.get("size", len(metadata["text"])),
                "distance": float(1.0 - score),
                "similarity": float(score),
            }
        )
        if len(retrieved) >= semantic_top_k:
            break

    return build_retrieval_payload(retrieved, strategy="semantic", query_type="general")


def retrieve_by_resume_intent(question, chunks):
    if not chunks:
        return None

    project_chunks = [chunk for chunk in chunks if chunk.get("section") == "Projects"]
    project_chunk = find_project_chunk(question, project_chunks)
    if project_chunk:
        return {
            "chunks": [project_chunk],
            "strategy": "project",
            "query_type": "specific_project",
        }

    if is_summary_query(question):
        selected_chunks = select_summary_chunks(chunks)
        if selected_chunks:
            return {
                "chunks": selected_chunks,
                "strategy": "summary",
                "query_type": "summary",
            }

    requested_section = detect_requested_section(question)
    if requested_section:
        section_chunks = [
            chunk
            for chunk in chunks
            if chunk.get("section") == requested_section
        ]
        if section_chunks:
            return {
                "chunks": section_chunks,
                "strategy": "section",
                "query_type": "section",
            }

    return None


def find_project_chunk(question, project_chunks):
    question_normalized = normalize_for_match(question)
    if not question_normalized:
        return None

    for chunk in project_chunks:
        title = chunk.get("title") or ""
        title_normalized = normalize_for_match(title)
        if title_normalized and title_normalized in question_normalized:
            return chunk

    subject = extract_project_subject(question)
    if subject:
        subject_normalized = normalize_for_match(subject)
        for chunk in project_chunks:
            title_normalized = normalize_for_match(chunk.get("title") or "")
            if subject_normalized and subject_normalized in title_normalized:
                return chunk

    return None


def extract_project_subject(question):
    match = re.search(
        r"\b(?:about|explain|describe|details of|detail of|tell me about)\s+([a-zA-Z0-9][a-zA-Z0-9 ._-]{1,80})",
        question,
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    return re.sub(r"[?.!,]+$", "", match.group(1)).strip()


def detect_requested_section(question):
    question_normalized = normalize_for_match(question)
    if not question_normalized:
        return None

    for section, aliases in RESUME_SECTIONS.items():
        for alias in aliases | {section.lower()}:
            alias_normalized = normalize_for_match(alias)
            if re.search(rf"\b{re.escape(alias_normalized)}\b", question_normalized):
                return section
    return None


def is_summary_query(question):
    question_normalized = normalize_for_match(question)
    return any(term in question_normalized for term in SUMMARY_QUERY_TERMS)


def select_summary_chunks(chunks):
    selected = []
    for section in SUMMARY_SECTION_PRIORITY:
        section_chunk = next((chunk for chunk in chunks if chunk.get("section") == section), None)
        if section_chunk:
            selected.append(section_chunk)
        if len(selected) >= 5:
            break
    return selected or chunks[:4]


def build_retrieval_payload(retrieved, strategy, query_type):
    context = "\n\n".join(
        (
            f"[{chunk['document_name']} | Chunk {chunk['chunk_number']} | "
            f"{chunk['section']} | {chunk.get('title') or chunk['section']} | Page {chunk.get('page') or 'N/A'}]\n"
            f"{chunk['content']}"
        )
        for chunk in retrieved
    )

    return {
        "chunks": retrieved,
        "context": context,
        "context_size": len(context),
        "chunk_count": len(retrieved),
        "strategy": strategy,
        "query_type": query_type,
    }


def normalize_for_match(value):
    clean = re.sub(r"[^a-zA-Z0-9\s]", " ", str(value).lower())
    return re.sub(r"\s+", " ", clean).strip()


def log_generated_chunks(document_name, extracted_text, chunks):
    has_projects = bool(re.search(r"\bprojects?\b", extracted_text, flags=re.IGNORECASE))
    print("\n" + "=" * 50)
    print(f"Resume parsed: {document_name}")
    print(f"Extracted text contains Projects heading/text: {has_projects}")
    print(f"Generated chunks: {len(chunks)}")
    for chunk in chunks:
        print(
            "Chunk metadata | "
            f"section={chunk.get('section')} | "
            f"title={chunk.get('title')} | "
            f"page={chunk.get('page') or 'N/A'} | "
            f"chunk_id={chunk.get('chunk_id')} | "
            f"size={chunk.get('size')}"
        )
    print("=" * 50)


def build_prompt(question, retrieval, chat_history=None):
    memory = format_recent_history(chat_history or [])
    memory_block = f"\nRECENT CONVERSATION:\n{memory}\n" if memory else ""

    return f"""
You are a Resume Intelligence Assistant.

Use ONLY the resume context and recent conversation shown here.
Never invent employers, dates, skills, projects, certifications, achievements, education, or contact details.
Never mix details from different projects.
Never combine unrelated resume sections into one answer unless the user asks for a summary or overview.
If information is unavailable in the resume context, clearly say that the resume does not provide that information.
Use recent conversation only to resolve follow-up references, not to add facts outside the resume.
Mention the section name and chunk number you relied on when useful.
{memory_block}
RESUME CONTEXT:
{retrieval["context"]}

QUESTION:
{question}
""".strip()


def format_recent_history(messages, limit=MAX_MEMORY_MESSAGES):
    recent = messages[-limit:]
    formatted = []

    for message in recent:
        role = message.get("role", "user").title()
        content = str(message.get("content", "")).strip()
        if content:
            formatted.append(f"{role}: {content}")

    return "\n".join(formatted)


def ask_llm(question, document_or_documents, chat_history=None, top_k=None, vector_store=None):
    retrieval = retrieve_relevant_chunks(
        question,
        document_or_documents,
        top_k=top_k,
        vector_store=vector_store,
    )
    prompt = build_prompt(question, retrieval, chat_history)
    print("\n" + "="*50)
    print("Prompt Size:", len(prompt))
    print("Retrieved Chunks:", retrieval["chunk_count"])
    print("Context Size:", retrieval["context_size"])
    print("="*50)
    response = requests.post(
        OLLAMA_URL,
        json={
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False,
        },
        timeout=300,
    )
    response.raise_for_status()

    return {
        "answer": response.json()["response"],
        "retrieval": retrieval,
        "prompt_size": len(prompt),
    }


def export_chat(messages):
    lines = ["Resume Intelligence Assistant - Chat Export", ""]

    for message in messages:
        timestamp = message.get("timestamp") or datetime.now().isoformat(timespec="seconds")
        role = message.get("role", "message").upper()
        lines.append(f"[{timestamp}] {role}")
        lines.append(str(message.get("content", "")))
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def _as_document_list(document_or_documents):
    if document_or_documents is None:
        return []
    if isinstance(document_or_documents, dict):
        return [document_or_documents]
    return list(document_or_documents)
