from __future__ import annotations

from io import BytesIO
from pathlib import Path
import textwrap
from uuid import UUID

try:
    from docx import Document
except ImportError:  # pragma: no cover - dependency is installed via requirements.txt in normal setup.
    Document = None

try:
    import fitz
except ImportError:  # pragma: no cover - dependency is installed via requirements.txt in normal setup.
    fitz = None


PREVIEW_CACHE_DIR = Path("backend/preview_cache")
PAGE_WIDTH = 595
PAGE_HEIGHT = 842
MARGIN = 44
FONT_SIZE = 10.5
LINE_HEIGHT = 14


class PreviewConversionError(RuntimeError):
    pass


def is_docx_file(file_name: str | None, mime_type: str | None) -> bool:
    value = f"{file_name or ''} {mime_type or ''}".lower()
    return ".docx" in value or "wordprocessingml.document" in value


def is_pdf_file(file_name: str | None, mime_type: str | None) -> bool:
    value = f"{file_name or ''} {mime_type or ''}".lower()
    return ".pdf" in value or "application/pdf" in value


def get_docx_preview_pdf(resume_id: UUID, original_file_name: str | None, resume_blob: bytes) -> bytes:
    PREVIEW_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cache_path = PREVIEW_CACHE_DIR / f"{resume_id}.pdf"

    if cache_path.exists() and cache_path.stat().st_size > 0:
        return cache_path.read_bytes()

    try:
        text = extract_docx_text(resume_blob)
        pdf_bytes = render_text_preview_to_pdf(
            title=Path(original_file_name or "resume.docx").name,
            text=text,
        )
        cache_path.write_bytes(pdf_bytes)
        return pdf_bytes
    except Exception as exc:
        cache_path.unlink(missing_ok=True)
        raise PreviewConversionError("DOCX preview conversion failed.") from exc


def extract_docx_text(resume_blob: bytes) -> str:
    if Document is None:
        raise PreviewConversionError("python-docx is not installed.")

    document = Document(BytesIO(resume_blob))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    if not text:
        raise PreviewConversionError("DOCX did not contain previewable text.")
    return text


def render_text_preview_to_pdf(title: str, text: str) -> bytes:
    if fitz is None:
        raise PreviewConversionError("PyMuPDF is not installed.")

    pdf = fitz.open()
    page = pdf.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)
    y = MARGIN

    def add_page():
        return pdf.new_page(width=PAGE_WIDTH, height=PAGE_HEIGHT)

    page.insert_text((MARGIN, y), title, fontsize=13, fontname="helv", color=(0.08, 0.14, 0.22))
    y += LINE_HEIGHT * 2

    for raw_line in text.splitlines():
        wrapped_lines = textwrap.wrap(raw_line, width=92) or [""]
        for line in wrapped_lines:
            if y > PAGE_HEIGHT - MARGIN:
                page = add_page()
                y = MARGIN
            page.insert_text((MARGIN, y), line, fontsize=FONT_SIZE, fontname="helv", color=(0.12, 0.16, 0.22))
            y += LINE_HEIGHT
        y += 4

    pdf_bytes = pdf.tobytes(garbage=4, deflate=True)
    pdf.close()
    return pdf_bytes
